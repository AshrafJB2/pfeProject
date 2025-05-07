from rest_framework import generics
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from django.http import FileResponse
from io import BytesIO
from docx import Document as DocxDocument
from .models import Content
from .serializers import ContentSerializer, UserSerializer, GetUser
from .utils import extract_text, process_with_gemini
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
import re
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated


class RegisterView(generics.CreateAPIView):
    serializer_class = UserSerializer
    permission_classes = [AllowAny]

class DownloadContentView(generics.RetrieveAPIView):
    permission_classes = [IsAuthenticated]
    queryset = Content.objects.all()
    serializer_class = ContentSerializer
    lookup_field = 'pk'

    def get(self, request, *args, **kwargs):
        # Retrieve the content object
        content = self.get_object()
        # Extract format from query parameters
        fmt = request.query_params.get('format', 'pdf').lower()
        # Sanitize filename
        filename = self.sanitize_filename(content.auto_title)

        # Dispatch to the correct generator
        if fmt == 'docx':
            return self.generate_docx(content, filename)
        elif fmt == 'pdf':
            return self.generate_pdf(content, filename)
        else:
            return self.generate_txt(content, filename)

    def sanitize_filename(self, filename):
        """Remove special characters and collapse spaces to underscores"""
        filename = re.sub(r'[^\w\s-]', '', filename or '').strip()
        return re.sub(r'[-\s]+', '_', filename)[:50]

    def generate_docx(self, content, filename):
        """Generate a .docx file response"""
        doc = DocxDocument()
        doc.add_heading(content.auto_title or 'Summary', level=1)
        if content.summary:
            for para in content.summary.split('\n'):
                doc.add_paragraph(para)
        if content.keywords:
            doc.add_paragraph()  # blank line
            doc.add_paragraph(f"Keywords: {content.keywords}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return FileResponse(
            buffer,
            as_attachment=True,
            filename=f"{filename}.docx",
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    def generate_pdf(self, content, filename):
        """Generate a .pdf file response"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'Title', parent=styles['Heading1'], fontSize=18,
            alignment=TA_CENTER, spaceAfter=20
        )
        body_style = ParagraphStyle(
            'Body', parent=styles['BodyText'], fontSize=12,
            leading=14, spaceAfter=12
        )
        keyword_style = ParagraphStyle(
            'Keywords', parent=styles['BodyText'], fontSize=12,
            textColor='#555555', spaceBefore=20
        )

        story = []
        story.append(Paragraph(content.auto_title or 'Summary', title_style))
        story.append(Spacer(1, 12))

        if content.summary:
            for para in content.summary.split('\n'):
                story.append(Paragraph(para, body_style))

        if content.keywords:
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>Keywords:</b> {content.keywords}", keyword_style))

        doc.build(story)
        buffer.seek(0)

        return FileResponse(
            buffer,
            as_attachment=True,
            filename=f"{filename}.pdf",
            content_type='application/pdf'
        )

    def generate_txt(self, content, filename):
        """Generate a plain .txt file response"""
        parts = [content.auto_title or 'Summary', '']
        if content.summary:
            parts.append(content.summary)
        if content.keywords:
            parts.extend(['', f"Keywords: {content.keywords}"])

        text_bytes = '\n'.join(parts).encode('utf-8')
        return FileResponse(
            BytesIO(text_bytes),
            as_attachment=True,
            filename=f"{filename}.txt",
            content_type='text/plain'
        )


class ContentCreateView(generics.ListCreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ContentSerializer

    def get_queryset(self):
        # Filter content by the authenticated user
        return Content.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        # Save the content with the authenticated user
        instance = serializer.save(user=self.request.user)

        # Extract text from file if provided
        if instance.original_file:
            instance.extracted_text = extract_text(instance.original_file, instance.original_file.name)
        else:
            instance.extracted_text = instance.original_text

        # Process with Gemini
        results = process_with_gemini(
            instance.extracted_text,
            instance.summary_length
        )

        instance.summary = results.get('summary', '')
        instance.keywords = results.get('keywords', '')
        instance.auto_title = results.get('title', 'Content Summary')
        instance.save()


class ContentDetailView(generics.RetrieveAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ContentSerializer

    def get_queryset(self):
        # Filter content by the authenticated user
        return Content.objects.filter(user=self.request.user)


class CurrentUserView(APIView):
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        serializer = GetUser(request.user)
        return Response(serializer.data)
